#!/usr/bin/env python3
"""
Extrae ejercicios desde BASE DE DATOS.docx y genera exerciseDatabaseExtended.json.

Requisitos: pip install python-docx

Uso:
  python scripts/extract_exercise_db_from_docx.py
  python scripts/extract_exercise_db_from_docx.py "ruta/al/archivo.docx"

El script busca tablas en el docx y mapea columnas a ExerciseMuscleInfo.
Si la estructura del docx difiere, edita COLUMN_MAP y la lógica de parseo.
"""

import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    print("Instala python-docx: pip install python-docx")
    sys.exit(1)

# Mapeo flexible: clave (debe estar contenida en header normalizado) -> campo ExerciseMuscleInfo
# Orden: las claves más específicas primero para evitar falsos positivos
COLUMN_MAP = [
    ("id", "id"),
    ("nombredelejercicio", "name"),
    ("nombreejercicio", "name"),
    ("ejercicio", "name"),
    ("nombre", "name"),
    ("name", "name"),
    ("alias", "alias"),
    ("descripcion", "description"),
    ("description", "description"),
    ("musculos", "involvedMuscles"),
    ("musculo", "subMuscleGroup"),
    ("submusculo", "subMuscleGroup"),
    ("categoria", "category"),
    ("category", "category"),
    ("tipo", "type"),
    ("type", "type"),
    ("equipo", "equipment"),
    ("equipment", "equipment"),
    ("patron", "force"),
    ("force", "force"),
    ("bodypart", "bodyPart"),
    ("cadena", "chain"),
    ("chain", "chain"),
    ("setup", "setupTime"),
    ("setuptime", "setupTime"),
    ("dificultad", "technicalDifficulty"),
    ("technicaldifficulty", "technicalDifficulty"),
    ("transferencia", "transferability"),
    ("transferability", "transferability"),
    ("riesgo", "injuryRisk"),
    ("injuryrisk", "injuryRisk"),
    ("efc", "efc"),
    ("cnc", "cnc"),
    ("ssc", "ssc"),
]


def slugify(name: str) -> str:
    """Genera id desde nombre: 'Press de Banca' -> 'press-de-banca'."""
    s = name.lower().strip()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[-\s]+", "-", s)
    return s[:50] if s else "ejercicio"


def parse_muscles(val: str) -> list:
    """Convierte texto de músculos a involvedMuscles."""
    if not val or not str(val).strip():
        return []
    parts = [p.strip() for p in str(val).replace(",", ";").split(";") if p.strip()]
    result = []
    for i, p in enumerate(parts[:6]):  # max 6 músculos
        role = "primary" if i == 0 else ("secondary" if i < 3 else "stabilizer")
        act = 1.0 if i == 0 else (0.6 if i == 1 else (0.5 if i == 2 else 0.3))
        result.append({"muscle": p, "role": role, "activation": act})
    return result


def safe_float(val, default=None):
    try:
        return float(str(val).replace(",", "."))
    except (ValueError, TypeError):
        return default


def safe_int(val, default=None):
    try:
        return int(float(str(val)))
    except (ValueError, TypeError):
        return default


def row_to_exercise(headers: list[str], row_cells: list) -> dict | None:
    """Convierte una fila de tabla a objeto ejercicio."""
    row = [c.text.strip() if hasattr(c, "text") else str(c) for c in row_cells]
    if not row:
        return None
    data = {}
    used_fields = set()
    for i, h in enumerate(headers):
        if i >= len(row):
            break
        hnorm = h.lower().replace(" ", "").replace("_", "").replace("-", "")
        if not hnorm:
            continue
        for col_key, field in COLUMN_MAP:
            if field in used_fields and field != "involvedMuscles":
                continue
            ck = col_key.replace(" ", "").replace("-", "")
            if ck in hnorm or hnorm in ck:
                v = row[i]
                if not v:
                    continue
                if field == "involvedMuscles":
                    data[field] = parse_muscles(v)
                elif field in ("setupTime", "technicalDifficulty", "transferability"):
                    data[field] = safe_int(v) or safe_float(v)
                elif field in ("efc", "cnc", "ssc"):
                    data[field] = safe_float(v)
                elif field == "injuryRisk":
                    if isinstance(v, (int, float)):
                        data[field] = {"level": int(v), "details": ""}
                    else:
                        data[field] = {"level": 5, "details": str(v)[:200]}
                else:
                    data[field] = v
                used_fields.add(field)
                break
    if "name" not in data:
        for i, cell in enumerate(row):
            if cell and len(cell) >= 3 and len(cell) <= 80 and not re.match(r"^[\d\s.,]+$", cell):
                data["name"] = cell
                break
    if "name" not in data:
        return None
    data.setdefault("id", f"ext_{slugify(data['name'])}")
    data.setdefault("involvedMuscles", parse_muscles(data.get("subMuscleGroup", "")))
    data.setdefault("category", "Hipertrofia")
    data.setdefault("type", "Accesorio")
    data.setdefault("equipment", "Otro")
    data.setdefault("force", "Otro")
    data.setdefault("bodyPart", "upper")
    data.setdefault("chain", "anterior")
    data.setdefault("efc", 2.5)
    data.setdefault("cnc", 2.5)
    data.setdefault("ssc", 0.2)
    return data


def extract_from_docx(path: Path) -> list[dict]:
    """Extrae ejercicios de un archivo .docx."""
    doc = Document(path)
    exercises = []
    seen_ids = set()
    id_counter = {}

    def make_unique_id(base_id: str) -> str:
        if base_id not in seen_ids:
            return base_id
        id_counter[base_id] = id_counter.get(base_id, 1) + 1
        return f"{base_id}_{id_counter[base_id]}"

    for table in doc.tables:
        if not table.rows or len(table.rows) < 2:
            continue
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        has_name_col = any(
            "nombre" in h or "name" in h or "ejercicio" in h or "exercise" in h
            for h in headers
        )
        if not has_name_col and len(headers) > 0:
            headers_flat = "".join(headers).replace(" ", "")
            if any(x in headers_flat for x in ["nombre", "name", "ejercicio", "exercise"]):
                has_name_col = True
        if not has_name_col:
            for row in table.rows[1:]:
                cells = [c.text.strip() if hasattr(c, "text") else str(c) for c in row.cells]
                for cell in cells:
                    if cell and len(cell) >= 4 and len(cell) <= 80 and not re.match(r"^[\d\s.,]+$", cell):
                        base_id = f"ext_{slugify(cell)}"
                        uid = make_unique_id(base_id)
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            exercises.append({
                                "id": uid,
                                "name": cell,
                                "involvedMuscles": [],
                                "category": "Hipertrofia",
                                "type": "Accesorio",
                                "equipment": "Otro",
                                "force": "Otro",
                                "bodyPart": "upper",
                                "chain": "anterior",
                                "efc": 2.5,
                                "cnc": 2.5,
                                "ssc": 0.2,
                            })
                        break
            continue
        for row in table.rows[1:]:
            ex = row_to_exercise(headers, row.cells)
            if ex:
                base_id = ex.get("id", f"ext_{slugify(ex['name'])}")
                ex["id"] = base_id
                uid = make_unique_id(base_id)
                if uid != base_id:
                    ex["id"] = uid
                if ex["id"] not in seen_ids:
                    seen_ids.add(ex["id"])
                    exercises.append(ex)

    # Si no hay tablas o pocos ejercicios, intentar párrafos y listas
    if len(exercises) < 100:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) < 4:
                continue
            name = None
            desc = ""
            if ":" in text:
                parts = text.split(":", 1)
                name = parts[0].strip()
                desc = parts[1].strip()[:500] if len(parts) > 1 else ""
            elif re.match(r"^[A-Za-zÁ-ú\s\-]+$", text) and 4 <= len(text) <= 80:
                name = text
            if name and len(name) > 3 and len(name) < 80 and not re.match(r"^[\d\s.,]+$", name):
                base_id = f"ext_{slugify(name)}"
                uid = make_unique_id(base_id)
                if uid not in seen_ids:
                    seen_ids.add(uid)
                    exercises.append({
                        "id": uid,
                        "name": name,
                        "description": desc,
                        "involvedMuscles": [],
                        "category": "Hipertrofia",
                        "type": "Accesorio",
                        "equipment": "Otro",
                        "force": "Otro",
                        "bodyPart": "upper",
                        "chain": "anterior",
                        "efc": 2.5,
                        "cnc": 2.5,
                        "ssc": 0.2,
                    })

    return exercises


def main():
    root = Path(__file__).resolve().parent.parent
    docx_path = root / "BASE DE DATOS.docx"
    if len(sys.argv) > 1:
        docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"No se encontró: {docx_path}")
        print("Coloca BASE DE DATOS.docx en la raíz del proyecto o pasa la ruta como argumento.")
        # Generar JSON vacío para que la estructura exista
        out_path = root / "data" / "exerciseDatabaseExtended.json"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump([], f, ensure_ascii=False, indent=2)
        print(f"Creado {out_path} vacío. Ejecuta el script cuando tengas el docx.")
        return

    exercises = extract_from_docx(docx_path)
    out_path = root / "data" / "exerciseDatabaseExtended.json"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(exercises, f, ensure_ascii=False, indent=2)

    print(f"Extraídos {len(exercises)} ejercicios -> {out_path}")


if __name__ == "__main__":
    main()
